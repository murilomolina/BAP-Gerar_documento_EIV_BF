import base64
import tkinter as tk
from tkinter import messagebox
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import locale
from docx.shared import Inches

imagem_base64 = "data:image/jpeg;base64,/9j/4AAQSkZJRgABAQEAYABgAAD/2wBDAAMCAgMCAgMDAwMEAwMEBQgFBQQEBQoHBwYIDAoMDAsKCwsNDhIQDQ4RDgsLEBYQERMUFRUVDA8XGBYUGBIUFRT/2wBDAQMEBAUEBQkFBQkUDQsNFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBT/wAARCABWAJsDASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwDyzSeVWuw0vIxXH6TxGD2AzUNn8bPA9q5SXX4UdTggxvwR+Fft8qc6l+SLZ8q0+h7NpLHA/wAK7DSW5XivDdO/aC+HkeA/iaBf+2Un/wATXrvgjxRpHi7T1vtF1K31O0ztMlu+dp9COoP1rz61GpTV5xaXoc8otbo9D03txXR2X3hXnPib4ieHPhrpUGo+JtVj0mzml8mOSRWbc+CcAKCegNO8B/tDfDrx34gtdD0HxRb6hq1znybZY5FL4GSASoHSvOdCrKHPGLa720OeUW1ex7DZ9BWzF2rjPE3jTRPh74duNe8RX6aZpFuVWW6kVmCljhRgAnkmuJj/AGz/AIKjr48tP/Aeb/4isI4etVXNTg2vJNmKhKWyPcY+1WR9w14ppv7YnwX1C6SCLx/p6SOcDzY5UX8SVwK9msbyDU7OC6s54ry1uEEkM8Dh0kU9CpHBFc9ajVov95Br1TRMoyj8SsNk71Wkryf4h/tdfCX4a6pLpur+LIZtRiJEttp0bXJjPcMygqD7ZrP8G/tkfCHx9qUWn6f4sjs72U7Y4tSha3Dk9AGYbc/jXRHB4lw5/Zu3ezKVOdr20PUrz+KsG+q/4u1/T/Cuh3+sarcrZ6ZZRefcXDAlUTj5uOSOR0rx2y/al+FXiLVLTTtP8ZWs95dyLDDGYpF3OTgDJXA/GnSo1ai5oRbS7IqMZPVI7LUu9clqzdcCtH4jeN9D+Hmkf2n4i1BdMsDKIPOkVmG85wMAE9jXC6L8UPC3j6w1G90DV49StdPG66kRGURjGe4GePSuynTm486i7d+h0RTtexDqjH5u1cfqeTnmsu//AGhfh1MDs8TwH/tlJ/8AE1zd/wDHLwJLnZ4hhb/tm/8AhXq08NX/AJH9zOiMZdi1qn8VfrL8DP8Akifw+/7F7T//AEmjr8j11e017TY7/T5xc2cwJjlAIDY4PWv1w+Bn/JE/h9/2L2n/APpNHXyvEiao00+7/I9fBfEz8j9L/wBQ3/XNv5V8habp9pq3jCGzvrsWFnPd+XNdN0iUty34V9e6X/qD/wBcz/KvkXR9Jg17xpbadc3Is7e6vfKkuGxiMFsFufSv07LNPau9v6ZxQ6nefEb4XeCfCvhdtQ0Tx1BreoCVUWxVRuZTnLcdMV6v+wJHf/8ACQeKnAk/sv7LGGJzsMu/jHbOM/hXn/xG+Anhrwb4NvdYsfG1rqN3bldlnlC0uWAwNrE55z+Fdn+wz8RLrS/FWqeGrqQHRZLSXUSCAPLeJcls+44p4mTrZfUcJOfqrdvImetN21H/ALe3jj+0/Gui+FoZMw6TbefMoPHnSevuFA/76rxLwXqWqfBv4reHtUu4ms7zTbqC6dG6+W2CfzRj+dQ+MvGFv46+LGoeINWeQ6fe6l5spjGWFuHAwB67AK639qL4ieEvil8RIde8IxXcFo9jDbzx3cQjIeNdgIAJ42ha68PRdGlTwrjdOLu/P/g3ZcY8qULH3x+25dQah+yr4hu7dg9vcPZTRMOhVpUYfoa+EP2W/wBnRP2j/FGsaQ/iCHw8NPsxdedNHvEnzY29RX0Z4g8e/wDCwP8Agm7JcSyebe6bLb6ZcEnnMc67SfqpH5V8h/Cf4W+OPipqt7Y+BtPutRvraHzp0tZhEVjzjJJYZGa8rLKcqGDrU+fkcZNXfTbuc9GLjTkr2szqP2m/2e/+GdfGVjoY8RWfiNbu0F0JbUbWi5xtcZOD3+lfQHhr40eJPhn/AME64TBczW+oarrU+j6bdZO+K1I3SFD26Oo9N1fJXxK+Hni34ZeJm0nxnpl1purmNZvLu3Ds6HowYEgj8a+tf2hrwePv2BfhT4g07SodLtNNvza3FrZqRFGQJI9+P9plyc92rpxSU4YaFVqacl72ltm195c9VBS113PEv2Vf2Zbz9pfxZqdvLqh0jR9NiE99fbPMkJY4VEB6sT3NaH7Wf7J9x+zXf6Pc2mrtrmgapuWG6kj8uWKVeSjgcdOQR6V69/wTH+JGheHtc8X+GNV1C20281SOKeze6kEayshO5Ax4zg5A71uf8FOviPoOqaT4R8JafqNtqOqW91Jf3K2sqyCBdhRVYqSATuJx7VhLGYtZssOv4fa3S2/3k+0qe35OhV+DvxU1D4jfsM/EzStXuHu77w7ZNaxzytudrdihQMe+05X6AV8HQSvbypLEzJJGQyuvBUg5Br65/Zi0W4g/ZH+PGrOrLbXFsltGx6MybWbH/fQrzz9kr4Z2Pxe1rxl4ZvVUPc6HI1rMwyYZ1dTGw/H+ddOHnTwbxM7e6pfmlcuDVPnfS56f8bPisnxg/Y50XV5JA2q2uqQWeornkTIjjf8A8CGG/E1kfscf8k3+JP8A1x/9ptXzjqN5rfg6217wfehrdGu1F5avn5ZoSwBH/fR57ivpD9jf/km3xK/64f8AtNqjE4eOGwU4w+FyTXo2glHlg0u58yeD9J0/XPEVlY6pqSaPYTMRLfSDKxDBOT+PH416bJ8IfAJlWOD4k2c7s21VWLqT0rzTwX4ei8V+JbHSp9Qh0qG5Yq15cfcjwCcnkelerx/ADSdMvYJx8QdFm8mRXwpHODn+/Xp4qooSSdRx02Sv+jNpPzPWNB8Mr4P8M2ejrObkWysBKVxuyc9K/X74Gf8AJE/h9/2L2n/+k0dfkpdX1tqVv59pcRXUB4EkLhlJHXkV+tfwM/5In8Pv+xe0/wD9Jo6/IOJm5UoOW93+R14L4mfkhpQ3R7R1ZCPzFeHt+y/4q1C6mkjudPCyOzDdN6nPpXuOj/dX6V2OjzRynCSxyFeuxw2Prg8V9tTxVXCtun1PN5nHY+a7f9jrxnc423elj6zn/CvXvh3+yrrngjwT4oNpqVnc+LtYszp8DBisNrC5/eHcRksRxXs+j3EUw/dzQyhfvbJFbb9cHiuv0W5huFLQzwyqv3mjkVgPqQeKwr5piqkeVtW9DGVWb0Plb4Y/sC6hLrzN431CEaN5LbRpc+ZTJkbeo4GM12fxE/4J46Tc+H4v+EE1G4GtCYbxq04EJiwc4IHBzivqLSbiGeHzY54ZIV+9IkqlB9WBwK6PTbq3lhMyXVu8C9ZlmQov1bOBXFUzfG86nz7dOhjKvUve58i+Cf2OviVovwP8eeAbm90d01ya0u7NluSVjliky+eOAV/UCu//AGLf2UvF/wCz34017VvEVzp09rf2AtYxZTF2Db888dK+mbG4gkt/PW5t3txyZlmUxj6tnA/OtaK4g8lJftNv5L/dk85djfRs4NcdbNMTVhOnK1pvXT0/yMZV5yTT6nyl+2f+yP4u/aG8daHrfhu802GCz0/7JMt9MUbcGJGOORivRvgf+zvc6H+zTcfCr4iQ2d9bSSTqxsZdw8t2Dq6t2dWyR9BXrUHjLRJPFf8AwjK6lC+u/ZBffY1yT5B/5abvu4/Gt23uYLqN2t7mC5VThjBMsm0+h2k4rlqY7Eewhh3pGNmtNfJ3IdWfKodEfmP8Qv8Agmf8QtB1eVvCGoaf4i03fmB5JxbXCr23BuM+4NUvB3/BNr4ma5qaHxNd6b4fsSwM0zXIuZSvfaq5yfqa/T67uYLXYZ7mC38w7U86VY9x9BkjJ+lRzfuwS5CAdSxAA/E16kc/x3Jy3V+9tf8AL8Db61VtY8S1z4E2Og/s56x8MvByJbiewe3hmujjzZmYM0khHckH6dK8L/ZX/ZT8XfArx1qOta/c6fNa3Fg1qos5i7bywPp04NfZV1NDIJNtxC4Tl9synaPU4PH41hXs0PlyOLiAxp9+QSrtX6nOB+NcNPHV4050r3U9X3M41ZpOPc+Ov2of2RdR+J/jmPxN4WmsrWe5iC6hFcvsDyLwsi8dSOD9Kd8CPgP4g+EvhHxhpesT2ctxq0ey3NtJuUHYR8x7cmvo/wAUeKtI0F9NS+v4oTqc/wBnsyMuJ5P7qlcjNY2q3EXmzRedEZYgS8YkUsgHUkZyK9COOxMqCoS+Hpp2OlVJuPK9j4Hn/Y28a24O+70v8Jz/AIVmXH7Kviy3+9c6cfpMf8K+zofEmleItMbUNN1G3vLDe0RuEcBN6nBXJxzmud1iWKJN7yxojdGZwAfoSea9iOa4vaVvuOlVZ9TyHwB4RvPBHg6PSr9o3uEmkkJhOVw2Mc1+yfwM/wCSJ/D7/sXtP/8ASaOvyb1XByQcg9CDkGv1k+Bn/JE/h9/2L2n/APpNHXw3E83UpwnLdt/kengtZSZ+QF7Z3epeHb20sJI4byaApG8udoJ9cc1yPhf4LeJbW58xru2t45dOls5o4ruQb3ZSFbIHHP1IrvtI+6v0rsdNbp619msROinGPU8zma2PJ/CHwC8XWltrEEt5ZQQXlgbRUW6k3u+4EHeBwOOuM1s+F/2bfHSeH/EmnjW7PSTqNrFFF5U8jtIyNnDMANqnoSBmvbtKPArr9J/hrnqZjX12+7sZSqSPBNI/Zf8AGc3hfxLZvfW1jJqMEEUVjDqcoid42BLsQvyng9Oueau6P+yh8Qv+EG8R6RPqGmxtfXVvc21rFfSocx53F5AMcg/dxj1r6e0wcCujsfvCuKWaYjXbe+3a3+RhKtI+VNJ/ZJ+JH/CB3GlT6vpLL/bcOqDQ2uphBexKpDxTSr0zkcAYrb0z9jfxfeaXpdlquoafb6O3iZtWuvD1pfTm3s7JkCmCFz8xPHTpwK+rrOtiLoK5pZpiX1X3f1/mZe3mfMfjD9lHxHa+JNbfwBqdtYaFL4X/ALCsINWu5ZZVcuWcMw5AI4DZyKn/AGb/ANmjxz8JfirH4l1e+sYdEk0U2FzptpeyTZuAVIkwVA7H1IzX09H2qwPuVyyzCu6bpO1mrbake2lblZ8o/tJfsxeNvih8UJvEujalaX+k3GmLYLpl9eyW32GQDBkTaCOeuRzmvQPFXgTX9F/ZT1fwel1ca74jg8PS2cc0UjPLPNtO0KxwxIyACeeK9nkqvIcUvrtWUKcJWtC1vkL2kmkux8IeDP2U/iFceGReSNp/hO4uvCp02S3ivbhpr66fBEl0G4Rl6Hb6cV0Hij9km4tPh7oFj4cuVXV4Z7a613Tb2/m+w6q0a4Zcj5lB/Wvrm95znmsG+rt/tLESlzXtqa+2m3c+JLj4I+OPC+peFdL+zSNDP4ik1dk0uZ5bLRoNu0Rh3OSR1x04FVND/Zs8WaN4qt5tS1S3ntIbueabUor2TzbqN8/IyY54PIJxX2LqTHaQDxXIar3rsjmNZ6aanQqsj5HX9njxHo+j6faKNLvbaxv57h9NkuZUgvY3ACFyOjLjpWHq/wADPEy6DpNt9vhu7i1aZmjN04jiDkkIuR8wHqa+pdU71yOpt1rtp4+u9Xb+v+HNo1JHmPhfQ77w14Ts9N1GaKe6hDAtDkqATwMnk1+wnwM/5In8Pv8AsXtP/wDSaOvya1X+Kv1l+Bn/ACRP4ff9i9p//pNHXx3E0nOnCT6t/kepgviZ8W2X/BPf4j2oUNrHhY49Ly5/+R63bP8AYV+INtjdqfhlvpeXH/yPX3bRXhPPsZLe33HV9TpHxfZfsbeO7XGb7w2fpeXH/wAYrds/2VfG1tjddeHz9Luf/wCMV9Z0Vi86xT3t9xH1Gi+h8z2v7OfjC3+9PoZ+l3N/8ZrVt/gR4rhIzJop/wC3ub/4zX0HXB3vxA1q58aahovh/wAOQ6va6Q9smqXU+pfZpEMwD7YIzGyyskZDtveIfMoBY5xKzLEzva39fMzlgcOt0zibf4N+KIRydHP/AG9y/wDxmr6fCvxKvbSf/AuX/wCM1haf+1lDfeC4dZPhloNQFtqFzcaa9+MxLb2bXcTK/l/Os0YUbto2sXHJQitzwj+0IfEtlYSnStPnF1rlvowu9F1f7dYN5sBl3x3HkpvZMbXj2jaf4q0licZHVxX9fM51hMHJ6X/EmX4Y+JF/h0n/AMC5f/jVS/8ACtvEe3GzSv8AwLl/+NVveJvidL4V8ZrpV7osqaR/ZF5qp1b7SmW+z+WXjSIZJ4lHzMV5GACORT8C/FTU/EUup2+teGhpV5b6Ta63bQafem+NzbT+btXmKPbMDCwKDcMsuHbJxj9cxPLz2Vv68y/qGFvy63+Zkt8MfEjfw6T/AOBcv/xqom+FfiRu2k/+Bcv/AMarN8D/ALR198RdN1STQ9N8KXd5p9pa6nOE8VPJaxWkyzE+bNHZuY54zDhoSmAGB3nodHQfjpq1/qnge11Hw/o9lJ4mgjuBptvrzT6pbRSJIyzm1Nsm6EbBukLrtycrldp1eJxcbppaf8P3IWDwbta+vqU7j4PeJ5s4/scf9vcv/wAZrNuPgX4rm6Poo/7e5v8A4zXZeLvjJdeG/ihpPhGDQ4LsXrWoEk180NzMsryLJJbQeSwnSBY98p8xNikZHK7srwP8cPEXxAj1mLSfB1jNqNkLbbB/bbLHbySSSrJb3rNbhre4iEW54kSYjenOGViLGYvl5rK39eZX1PCKXLrf5nHXX7Ovi+4zibQx9bub/wCM1i3n7LXjW5ztuvD4+t3P/wDGK9N1r4+SWPgPwx4kj0nym1CwudYvLNj5hS0t4SzmF8oG3yvbRxyMACs6uQPu1DJ8cvFC+IJ/CaeCLGTxtCr3LWC6432JrVYUkLrcfZt5k3SpHs8nGcnfjmtY47GrZL/hvmV9Wwq7/ieN3n7Hfju5zi98ODP/AE+XH/xisK8/Yb+IFxnbqXhpfreXH/yPX0F4X/aIm8XNo2pWHhhn8L6reJptvefb1+2/aWszcgG3KbBHx5e4zA7v4QvzV3vw38ZXXjrw7NqF7pX9i3cN9dWMtl9pFwUaGZ4jlwoBJ2ZwMgZwCepuWa46krySX9eppDDYeWkbnxJef8E/fiNcZ26v4XH1vLn/AOR6+4vhv4dufB/w78LaDevFJeaXpVrYzvbkmNpIoVRipIBK5U4yAcdhXmHiL9ojWPCNrfprnhjSNHv7a70+3JvvEWyyhW7EpX7Rc/ZyI5IxF86qrqA6EOQc17B4b1Rtc8P6bqLNZOby3juN2m3X2q2YMoYGKbavmIQch9oyCDgVx47GYrFQiq9rLa3odNGNKLap7mlRRRXinWFFFFABXF6l8L4LzxdPr9lr+taJJem3/tGz06aJYb/ySdm8vG0iHadpMTxllwCTgUUVUZOOxLipbnLXv7MPhK9s9JgN1qsTafol1oAminjDz200UkeZf3e1njE0xQgDBlfIIOK7XW/ANtr2peGbi41G+W10Gf7VDp8ZiENxMsbJG8pMZkygYkBHUEn5gw4oorR1qkt2SqcFsjO8VfCmHxd4ttNbvfEOrpb29rNZHR4VtfsksEwQTo5aAy4cRpkiQEY+Urk5r+E/g+ng+Qz23irXru9MVtafbLz7I8v2SBZBFa8W4GwGVnLY8xmxucjiiil7WfLy30D2cb83Ujf4K2N9o/iKz1fxBrmuXXiBLe21HUrySBJ5bSFiVtAsUKRrCweZWCoGInk+bJBGp4k+HI8V+ItO1HUPEGrPp1hdQ3sWhotqtp58RzHIW8jz8hsNgSgEjBGMglFHtZ3vcPZx2sYWpfAew1jW01K+8Sa7cNI9jNfQ5tEW/ms5TLbyyMluHjZW28QtGpC9OWLZMP7MelW+gz6VF4s8SxRPp9tpEM0bWSPb2ELl1tVC2wV42Jw/mq7MoKlsM4Yoq1iKi0TJ9jB62Ne3+Dp1KTxBaeJNRk13TrzRk0K1mk8uO5W3JkeZiIYo443LSRqPLXgW8R+9mqlx+zzY3Er6gfF3iVPE0jOJfEavaC8kgeFYWtyv2fyfL2oh4iDbkDbs5yUU3WnF6P8AqwlTg1qi5o/wD0Lw7qlrNpWo6pYaVZz/AG2z0NHheztroWpthOm+JpCwjJwrSFNx3bN3NWfDPwR0bQ4SNUu5/GE8eoSarZ3XiCzsZZbC6kkMjywGO3TYxc7s4OCBjAGKKKh1qj3ZSpwWyK/hz4M33hXSLuw0/wCI3itDc3H2p7qWLS5JzISS7M7WR3lsqCX3EBFAIAxXZeD/AArZeB/DOn6Fpxla0so9ivO++RySSzsePmZiScADJ4AGBRRUyqSnuVGEY7H/2Q=="
# Decodificar a imagem base64
imagem_decodificada = base64.b64decode(imagem_base64.split(",")[1])

# Caminho para a imagem temporária
caminho_imagem_temporaria = "imagem_logo.jpg"

# Salvar a imagem temporária no disco
with open(caminho_imagem_temporaria, "wb") as imagem_temporaria:
    imagem_temporaria.write(imagem_decodificada)

def formatar_num(numero):
    # Definir a localização para o formato monetário desejado
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

    # Formatando o número
    numero_formatado = locale.currency(numero, grouping=True, symbol=True)

    return numero_formatado

def atualiza_data():
    data_atual = datetime.now().date().strftime("%d-%m-%y")
    return data_atual

def hora_exata():
    hora_atual = datetime.now().strftime("%H:%M:%S")
    return hora_atual
nova_data = atualiza_data()
nova_hora = hora_exata()

def formata_data(nova_data):
    nova_data = nova_data.split('-')
    dic = {1: 'Janeiro',2:'Fevereiro',3: 'Março',4:"Abril", 5:"Maio", 6:"Junho", 7: 'Julho', 8:"Agosto", 9:"Setembro", 10:"Outubro", 11:"Novembro", 12:"Dezembro"}
    dia = (nova_data[0])
    mes = int(nova_data[1])
    ano = int(nova_data[2])
    if mes in dic:
        for chave, valor in dic.items():
            if chave == mes:
                return (f'{dia} de {valor} de 20{ano}')

def data_indice_CUB(nova_data):
    nova_data = nova_data.split('-')
    dic = {1: 'Janeiro',2:'Fevereiro',3: 'Março',4:"Abril", 5:"Maio", 6:"Junho", 7: 'Julho', 8:"Agosto", 9:"Setembro", 10:"Outubro", 11:"Novembro", 12:"Dezembro"}
    mes = int(nova_data[1])
    ano = int(nova_data[2])
    if mes in dic:
        for chave, valor in dic.items():
            if chave == mes:
                return (f'{valor}/{ano}')

def formatar_variaveis_negrito(texto):
    # Criar um documento temporário para acessar as funcionalidades do docx
    documento_temporario = Document()

    # Adicionar um parágrafo ao documento temporário
    paragrafo_temporario = documento_temporario.add_paragraph(texto)

    # Selecionar as variáveis entre {} e torná-las em negrito
    for variavel in paragrafo_temporario.runs:
        texto_variavel = variavel.text
        if texto_variavel.startswith('{') and texto_variavel.endswith('}'):
            variavel.bold = True

    # Obter o texto formatado com as variáveis em negrito
    texto_formatado = ""
    for run in paragrafo_temporario.runs:
        texto_formatado += run.text

    return texto_formatado

def atribuir_valor_m(): #adicionar o valor "m" para a variavel indice_CUB
    indice_CUB_entry.delete(0, tk.END)
    indice_CUB_entry.insert(0, 'm')

def toggle_campo_fmp():
    if fmp_entry.winfo_ismapped():
        fmp_entry.grid_remove()
    else:
        fmp_entry.grid(row=1, column=2, padx=5, pady=5)

def toggle_campo_CUB():
    if indice_CUB_entry.winfo_ismapped():
        indice_CUB_entry.grid_remove()
        nome_indice_label.grid_remove()
        nome_indice_CUB_entry.grid_remove()
    else:
        indice_CUB_entry.grid()
        nome_indice_label.grid()
        nome_indice_CUB_entry.grid()

def limpar_campos():
    nome_projeto_entry.delete(0, 'end')
    area_terreno_entry.delete(0, 'end')
    area_computavel_entry.delete(0, 'end')
    valor_referencia_entry.delete(0, 'end')
    zona_entry.delete(0, 'end')
    area_a_construir_entry.delete(0, 'end')
    fmp_entry.delete(0, 'end')
    indice_CUB_entry.delete(0, 'end')
    fmp_entry.delete(0, 'end')
    nome_indice_CUB_entry.delete(0, 'end')
    

def gerar_documento():
    nome_projeto = nome_projeto_entry.get()
    area_terreno = float(area_terreno_entry.get().replace(',', '.'))
    area_computavel = float(area_computavel_entry.get().replace(',', '.'))
    valor_referencia = float(valor_referencia_entry.get().replace(',', '.'))
    zona = int(zona_entry.get())
    area_a_construir = float(area_a_construir_entry.get().replace(',','.'))
    fmp = fmp_entry.get()
    nome_indice_CUB = nome_indice_CUB_entry.get()

    indice_CUB = indice_CUB_entry.get()
    if indice_CUB == '':
        indice_CUB = 1954.65
    else:
        try:
            indice_CUB = float(indice_CUB.replace(',', '.'))
        except ValueError:
            messagebox.showinfo("Valor inválido")

    if nome_indice_CUB == '':
        nome_indice_CUB = 'R8N'
    else:
        try:
            nome_indice_CUB = (nome_indice_CUB).upper()
        except ValueError:
            messagebox.showinfo("Valor inválido")

    if fmp == '':
        fmp = 5.3016
    else:
        try:
            fmp = float(fmp.replace(',','.'))
        except ValueError:
            messagebox.showinfo("Valor inválido")

    # Função para calcular a fórmula desejada
    #varaiveis referentes ao processamento:
    cp = area_computavel/area_terreno
    fator_reducao = 0.8
    if zona == 1:
        ic = 0.4
        coeficiente_basico = 2.5
        cpc = round(cp - coeficiente_basico,2)
        bf = (area_terreno * valor_referencia * cpc * ic * fator_reducao) * fmp
        dados_bf = f'''
            ÁREA DO TERRENO = {('{:,.2f}'.format(area_terreno).replace(',', '.'))} m²
            ÁREA COMPUTAVEL = {('{:,.2f}'.format(area_computavel).replace(',', '.'))} m²
            VALOR DE REFERENCIA = {(valor_referencia):.2f} FMP/m²
            COEFICIENTE BASICO = {(coeficiente_basico):.2f}
            COEFICIENTE PROJETO = {(cp):.2f}
            COEFICINTE PRETENDIDO = {(cp):.2f} - {(coeficiente_basico):.2f} = {(cpc):.2f} = Cp
            FMP = R${(fmp)}

            Contrapartida financeira:
            Bf = At x Vr x Cp x Ic x Fr
            Bf = {'{:,.2f}'.format(area_terreno).replace(',', '.')} x {valor_referencia:.2f} x {cp:.2f} x {ic:.2f} x {fator_reducao:.2f} x {fmp} = 
        '''
        if cp > coeficiente_basico :
            resultado_bf = f'RESULTADO OBTIDO:  {formatar_num(bf)};  CPC = {(cp):.2f}'
    elif zona == 2:
        ic = 0.33
        coeficiente_projeto = 3.0
        cpc = round(cp - coeficiente_projeto,2)
        bf = (area_terreno * valor_referencia * cpc * ic * fator_reducao) * fmp
        dados_bf = f'''
            ÁREA DO TERRENO = {('{:,.2f}'.format(area_terreno).replace(',', '.'))} m²
            ÁREA COMPUTAVEL = {('{:,.2f}'.format(area_computavel).replace(',', '.'))} m²
            VALOR DE REFERENCIA = {(valor_referencia):.2f} FMP/m²
            COEFICIENTE BASICO = {(cp):.2f}
            COEFICIENTE PROJETO = {(coeficiente_projeto):.2f}
            COEFICINTE PRETENDIDO = {(cp):.2f} - {(coeficiente_projeto):.2f} = {(cpc):.2f} = Cp
            FMP = R${(fmp)}

            Contrapartida financeira:
            Bf = At x Vr x Cp x Ic x Fr
            Bf = {'{:,.2f}'.format(area_terreno).replace(',', '.')} x {valor_referencia:.2f} x {cp:.2f} x {ic:.2f} x {fator_reducao:.2f} x {fmp}=
        '''
        if cp > coeficiente_projeto:
            resultado_bf = f'RESULTADO OBTIDO:  {formatar_num(bf)};  CPC = {(cp):.2f}'

    elif zona > 2 or zona < 1:
        messagebox.showinfo('\nERRO! Zona invalida!!\n') 

    if cp <= coeficiente_basico or cp <= 2.5:
        resultado_bf = f'Resultado do CP menor que 2.5, Não é necessario pagar a ODC; CPC = {(cp):.2f}'
        # Verificar a resposta 
        messagebox.showinfo("Título da Mensagem", resultado_bf)
            
    


    #tipo 1
    taxa_eiv = 0.025
    nr_vr = valor_referencia
    # float(input('Qual o valor do multiplicador referente ao valor de referencia ( nº x FMP / m² = VR )\nQual o valor de nº?\n').replace(',','.'))
    vr_eiv = round(fmp * nr_vr, 2)
        
    #calculo:
    eiv = ((area_terreno * vr_eiv) + ((area_a_construir) * (indice_CUB))) * taxa_eiv
    dados_eiv = f'''
        ÁREA DO TERRENO(At) = {'{:,.2f}'.format(area_terreno).replace(',', '.')} m²
        ÁREA Á CONSTRUIR(Ac) = {'{:,.2f}'.format(area_a_construir).replace(',', '.')}m²
        VR = VALOR DE REFERENCIA = {nr_vr} FMP/m² = R$ {vr_eiv}
        TAXA EIV/RIT - TIPO 1 = 2,5%
        Índice – {nome_indice_CUB} ({data_indice_CUB(nova_data)})  – CUB / SINDUSCON = R$ {indice_CUB}
        FMP = R$ {fmp}

        Calculo:
        EIV = [(At x VR) + (Ac x I_CUB_R8N)] x 2,5%
        EIV = ({(round(area_terreno, 2))} * {(round(vr_eiv, 2))} + {(round(area_a_construir,2))} * {(round(indice_CUB,2))}) x 0,025
        EIV = [({(round(area_terreno * vr_eiv, 2))}) + ({(round(area_a_construir * indice_CUB, 2))})] x 0,025
    '''

    # Criar um doc
    documento = Document()

    # Configurar cabeçalho
    secao = documento.sections[0]
    header = secao.header

    # Adicionar o parágrafo no cabeçalho
    paragrafo = header.add_paragraph()

    # Adicionar o "Texto 1" na esquerda
    texto1 = paragrafo.add_run("MG Barone Jr\nEngº Civil\n\n\n\n\t\t\t\t\t\t\t\t\t      ")
    texto1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph_format = paragrafo.paragraph_format
    paragraph_format.line_spacing = Pt(10)
    # Adicionar a imagem na direita
    imagem_path = caminho_imagem_temporaria
    largura_maxima = Inches(1.7)
    altura_maxima = Inches(0.89)

    run_imagem = paragrafo.add_run()
    run_imagem.add_picture(imagem_path, width=largura_maxima, height=altura_maxima)
    run_imagem.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


    header.add_paragraph("__________________________________________________________________________________________________________________")
    # Definir o espaçamento do cabeçalho (em polegadas)
    secao.header_distance = Inches(0.3)

    # Definir as margens para a seção principal
    secao = documento.sections[0]
    secao.left_margin = Inches(1)
    secao.right_margin = Inches(1)
    secao.top_margin = Inches(1)
    secao.bottom_margin = Inches(1)

    # titulo
    titulo_style = documento.styles.add_style("Titulo", WD_PARAGRAPH_ALIGNMENT.CENTER)
    fonte_titulo = titulo_style.font
    fonte_titulo.name = "Calibri"
    fonte_titulo.size = Pt(12)
    fonte_titulo.bold = True
    fonte_titulo.color.rgb = RGBColor(00, 00,00)  # Cor preta
    
    # Texto comum
    texto_style = documento.styles.add_style("Texto", WD_PARAGRAPH_ALIGNMENT.CENTER)
    fonte_texto = texto_style.font
    fonte_texto.name = "Calibri"
    fonte_texto.size = Pt(12)
    fonte_texto.bold = False
    fonte_texto.color.rgb = RGBColor(00,00,00)

    #texto Rodapé
    texto_rodape = documento.styles.add_style("Rodape", WD_PARAGRAPH_ALIGNMENT.CENTER)
    fonte_rodape = texto_rodape.font
    fonte_rodape.name = "Calibri"
    fonte_rodape.size = Pt(9)
    fonte_rodape.bold = False
    fonte_rodape.color.rgb = RGBColor(0,127,255)
    
    # Texto Negrito
    texto_style = documento.styles.add_style("Negrito", WD_PARAGRAPH_ALIGNMENT.CENTER)
    fonte_texto = texto_style.font
    fonte_texto.name = "Calibri"
    fonte_texto.size = Pt(12)
    fonte_texto.bold = True
    fonte_texto.color.rgb = RGBColor(00,00,00)

    
    # Adiciona um estilo de sombra ao documento
    sombra_style = documento.styles.add_style("Sombra", WD_PARAGRAPH_ALIGNMENT.CENTER)
    fonte_sombra = sombra_style.font
    fonte_sombra.name = "Calibri"
    fonte_sombra.size = Pt(12)
    fonte_sombra.bold = True
    fonte_sombra.color.rgb = RGBColor(0, 0, 0)


    #corrigindo formatação em negrito:
    dados_bf_negrito = formatar_variaveis_negrito(dados_bf)
    dados_eiv_negrito = formatar_variaveis_negrito(dados_eiv)



    # add titulo
    titulo = documento.add_paragraph(f"MEMORIAL DE CALCULOS BASICOS PARA OODC E EIV/RIT-TIPO I-LEI 9.924/16 PROJETO {nome_projeto.upper()}", style = "Titulo") 
    # Alinhe o parágrafo no centro
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # add paragrafos
    titulo_bf = documento.add_paragraph('1. CALCULO DO BENEFICIO FINANCEIRO - OUTORGA ONEROSA DO DIREITO DE CONSTRUIR:', style = "Negrito")
    texto_bf = documento.add_paragraph(f"{(dados_bf_negrito)}", style = "Texto")
    paragrafo = documento.add_paragraph(f"\t{resultado_bf}")
    paragrafo.add_run().bold = True
    paragrafo.style = sombra_style
    # Define a cor de fundo do parágrafo como laranja escuro
    paragrafo_para = paragrafo._element.get_or_add_pPr()
    shading_element = parse_xml(r'<w:shd {} w:fill="FFA500"/>'.format(nsdecls('w')))
    paragrafo_para.append(shading_element)
    #espaçamentos
    paragraph_format = titulo_bf.paragraph_format
    paragraph_format.line_spacing = Pt(12)
    paragraph_format = texto_bf.paragraph_format
    paragraph_format.line_spacing = Pt(12)

    # paragrafos
    titulo_eiv = documento.add_paragraph('2. CALCULO DO ESTUDO DE IMPACTO DE VIZINHANÇA E DE TRANSITO - TIPO 1:', style = "Negrito")
    texto_eiv = documento.add_paragraph(f"{(dados_eiv_negrito)}", style = "Texto")
    paragrafo = documento.add_paragraph(f"\tEIV =  {formatar_num(eiv)}")
    paragrafo.add_run().bold = True
    paragrafo.style = sombra_style
    # Define a cor de fundo do parágrafo como laranja escuro
    paragrafo_para = paragrafo._element.get_or_add_pPr()
    shading_element = parse_xml(r'<w:shd {} w:fill="FFA500"/>'.format(nsdecls('w')))
    paragrafo_para.append(shading_element)
    #espaçamentos
    paragraph_format = titulo_eiv.paragraph_format
    paragraph_format.line_spacing = Pt(12)
    paragraph_format = texto_eiv.paragraph_format
    paragraph_format.line_spacing = Pt(12)


    # espaço para assinatura
    data_lugar = documento.add_paragraph(f'Santo André, SP, {formata_data(nova_data)}\n')
    linha_para_assinatura = documento.add_paragraph('__________________________________')
    nome_assinatura = documento.add_paragraph('Miguel G. Barone Jr.\nEngº Civil')
    # alinhamento e espaçamentos
    paragraph_format = nome_assinatura.paragraph_format
    paragraph_format.line_spacing = Pt(12)
    data_lugar.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_para_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nome_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Rodapé
    footer = secao.footer
    paragrafo_rodape = footer.paragraphs[0]
    paragrafo_rodape.text = "________________________________________________________________________________________________________\n\t\t\tAVENIDA PORTUGAL, 397 – 4º ANDAR – SALA 406 – CENTRO – SANTO ANDRÉ – SP\n\t\tFONE: 11 - 4438.23.52 – Cel / WhatsApp : +55 11 - 9.9949.11.65 – e-mail: barone@terra.com.br"
    paragrafo_rodape.style = texto_rodape
    paragraph_format = paragrafo_rodape.paragraph_format
    paragraph_format.line_spacing = Pt(12)

        
    nome_arquivo = f"{nome_projeto}-BF_EIV-{nova_data}.docx"

    # Concatenar o caminho absoluto com o nome do arquivo
    caminho_completo = f"{nome_arquivo}"

    # print(f"Arquivo '{nome_projeto}-BF_EIV-{nova_data}/{nova_hora}.docx' foi criado com sucesso!")
    try:
        documento.save(caminho_completo)
    except Exception as erro:
        # Imprimir o possível erro
        messagebox.showinfo("ERRO CRITICO","Ocorreu um erro ao gerar o DOC:", erro)    

    messagebox.showinfo("Documento Gerado!",f"Arquivo '{nome_projeto}-BF_EIV-{nova_data}/{nova_hora}.docx' foi criado com sucesso!")

# Interface gráfica
root = tk.Tk()
root.title("Gerador de Documento")

# Labels e Entradas
tk.Label(root, text="Nome do Projeto:").grid(row=0, column=0, sticky="e", padx=5, pady=5)
nome_projeto_entry = tk.Entry(root)
nome_projeto_entry.grid(row=0, column=1, padx=5, pady=5)

tk.Label(root, text="Área do Terreno (m²):").grid(row=1, column=0, sticky="e", padx=5, pady=5)
area_terreno_entry = tk.Entry(root)
area_terreno_entry.grid(row=1, column=1, padx=5, pady=5)

tk.Label(root, text="Área Computável (m²):").grid(row=2, column=0, sticky="e", padx=5, pady=5)
area_computavel_entry = tk.Entry(root)
area_computavel_entry.grid(row=2, column=1, padx=5, pady=5)

tk.Label(root, text="Valor Referência (FMP/m²):").grid(row=3, column=0, sticky="e", padx=5, pady=5)
valor_referencia_entry = tk.Entry(root)
valor_referencia_entry.grid(row=3, column=1, padx=5, pady=5)

tk.Label(root, text="Zona (1 ou 2):").grid(row=4, column=0, sticky="e", padx=5, pady=5)
zona_entry = tk.Entry(root)
zona_entry.grid(row=4, column=1, padx=5, pady=5)

tk.Label(root, text="área a construir:").grid(row=5, column=0, sticky="e", padx=5, pady=5)
area_a_construir_entry = tk.Entry(root)
area_a_construir_entry.grid(row=5, column=1, padx=5, pady=5)

toggle_button_CUB = tk.Button(root, text="Alterar o índice\nCUB/SINDUSCON?\nAtual = R$ 1954.65", command=toggle_campo_CUB, bg="dark grey", bd=4, relief=tk.RAISED)
toggle_button_CUB.grid(row=3, column=2, padx=5, pady=5)


# Configurar coluna para expansão
root.columnconfigure(4, weight=1)

# Campo de preenchimento inicialmente oculto
indice_CUB_entry = tk.Entry(root)
indice_CUB_entry.grid(row=4, column=2, padx=5, pady=5)
indice_CUB_entry.grid_remove()

# Rótulo para o nome do índice
nome_indice_label = tk.Label(root, text='Alterar nome do Índice (Índice padrão: R8N)')
nome_indice_label.grid(row=5, column=2, padx=5, pady=5)
nome_indice_label.grid_remove()

# Campo para o nome do índice
nome_indice_CUB_entry = tk.Entry(root)
nome_indice_CUB_entry.grid(row=6, column=2, padx=5, pady=5)
nome_indice_CUB_entry.grid_remove()


# Botão para alternar o campo de preenchimento
toggle_button_fmp = tk.Button(root, text="Alterar valor do FMP\nAtual R$5,3016", command=toggle_campo_fmp, bg="dark grey",  bd=4, relief=tk.RAISED)
toggle_button_fmp.grid(row=0, column=2, columnspan=2, pady=10)

# Configurar coluna para expansão
root.columnconfigure(0, weight=1)

# Campo de preenchimento inicialmente oculto
fmp_entry = tk.Entry(root)
fmp_entry.grid(row=2, column=2, padx=5, pady=5)
fmp_entry.grid_remove()


# Botão para gerar o documento
tk.Button(root, text="Gerar Documento", command=gerar_documento, bg="light green", fg="black", bd=4, relief=tk.RAISED).grid(row=7, column=0, columnspan=2, pady=10)

#botão para limpar os campos
botao_limpar = tk.Button(root, text="Limpar Campos", command=limpar_campos, bg="red", fg="black", bd=4, relief=tk.RAISED)
botao_limpar.grid(row = 8, column=0, columnspan=2, pady=10)

# Iniciar a interface gráfica
root.mainloop()